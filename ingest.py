from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import chromadb
import fitz
from chromadb.api.types import EmbeddingFunction
from docx import Document
from openpyxl import load_workbook


SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xlsm", ".xls", ".docx", ".txt", ".md"}
COLLECTION_NAME = "quality_system_docs"


class ChineseEmbeddingFunction(EmbeddingFunction):
    def __init__(self, model_name: str = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2") -> None:
        self.model_name = model_name
        self.model = None
        self.fallback = HashEmbeddingFunction()

    def __call__(self, input: list[str]) -> list[list[float]]:
        try:
            if self.model is None:
                from sentence_transformers import SentenceTransformer

                self.model = SentenceTransformer(self.model_name)
            vectors = self.model.encode(input, normalize_embeddings=True)
            return [vector.tolist() for vector in vectors]
        except Exception:
            return self.fallback(input)


class HashEmbeddingFunction(EmbeddingFunction):
    """Small local embedding function so the MVP runs without API keys."""

    def __init__(self, dimensions: int = 384) -> None:
        self.dimensions = dimensions

    def __call__(self, input: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in input]

    def _embed(self, text: str) -> list[float]:
        vector = [0.0] * self.dimensions
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
        for token in tokens:
            digest = hashlib.sha256(token.encode("utf-8")).digest()
            index = int.from_bytes(digest[:4], "big") % self.dimensions
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vector[index] += sign
        length = sum(value * value for value in vector) ** 0.5
        if not length:
            return vector
        return [value / length for value in vector]


@dataclass
class Chunk:
    text: str
    metadata: dict[str, str | int]


def get_collection(reset: bool = False, db_dir: str | Path = "db"):
    client = chromadb.PersistentClient(path=str(db_dir))
    if reset:
        try:
            client.delete_collection(COLLECTION_NAME)
        except Exception:
            pass
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=ChineseEmbeddingFunction(),
        metadata={"hnsw:space": "cosine"},
    )


def ingest_folder(folder: str | Path = "docs", reset: bool = True, db_dir: str | Path = "db") -> int:
    folder_path = Path(folder)
    folder_path.mkdir(exist_ok=True)
    collection = get_collection(reset=reset, db_dir=db_dir)
    chunks = list(iter_chunks(folder_path))
    if not chunks:
        return 0

    for start in range(0, len(chunks), 1000):
        batch = chunks[start : start + 1000]
        collection.add(
            ids=[make_chunk_id(chunk) for chunk in batch],
            documents=[chunk.text for chunk in batch],
            metadatas=[chunk.metadata for chunk in batch],
        )
    return len(chunks)


def iter_chunks(folder: Path) -> Iterable[Chunk]:
    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            yield from split_records(extract_records(path))
        except Exception as exc:
            yield Chunk(
                f"文件解析失败：{exc}",
                base_metadata(path, path.suffix.lower().lstrip(".")) | {"status": "解析失败"},
            )


def extract_records(path: Path) -> Iterable[Chunk]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        yield from extract_pdf(path)
    elif suffix in {".xlsx", ".xlsm"}:
        yield from extract_excel(path)
    elif suffix == ".xls":
        yield from extract_legacy_excel(path)
    elif suffix == ".docx":
        yield from extract_docx(path)
    elif suffix in {".txt", ".md"}:
        yield from extract_text(path)


def extract_pdf(path: Path) -> Iterable[Chunk]:
    with fitz.open(path) as doc:
        for page_number, page in enumerate(doc, start=1):
            text = clean_text(page.get_text("text"))
            if text:
                yield Chunk(text, base_metadata(path, "pdf") | {"page": page_number})


def extract_excel(path: Path) -> Iterable[Chunk]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    for sheet in workbook.worksheets:
        rows: list[str] = []
        row_start = 0
        last_row_number = 0
        for row in sheet.iter_rows():
            if not row:
                continue
            last_row_number = row[0].row
            values = [format_cell(cell.value) for cell in row if format_cell(cell.value)]
            if values:
                if not rows:
                    row_start = last_row_number
                rows.append(" | ".join(values))
            if len(rows) >= 20:
                yield excel_chunk(path, sheet.title, row_start, last_row_number, rows)
                rows = []
        if rows:
            yield excel_chunk(path, sheet.title, row_start, last_row_number or row_start + len(rows) - 1, rows)


def extract_legacy_excel(path: Path) -> Iterable[Chunk]:
    import xlrd

    workbook = xlrd.open_workbook(path, on_demand=True)
    for sheet in workbook.sheets():
        rows: list[str] = []
        row_start = 0
        last_row_number = 0
        for row_index in range(sheet.nrows):
            last_row_number = row_index + 1
            values = [format_cell(sheet.cell_value(row_index, col)) for col in range(sheet.ncols)]
            values = [value for value in values if value]
            if values:
                if not rows:
                    row_start = last_row_number
                rows.append(" | ".join(values))
            if len(rows) >= 20:
                yield excel_chunk(path, sheet.name, row_start, last_row_number, rows)
                rows = []
        if rows:
            yield excel_chunk(path, sheet.name, row_start, last_row_number, rows)


def extract_docx(path: Path) -> Iterable[Chunk]:
    document = Document(path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    text = clean_text("\n".join(lines))
    if text:
        yield Chunk(text, base_metadata(path, "docx"))


def extract_text(path: Path) -> Iterable[Chunk]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    text = clean_text(text)
    if text:
        yield Chunk(text, base_metadata(path, path.suffix.lower().lstrip(".")))


def split_records(records: Iterable[Chunk], max_chars: int = 1200, overlap: int = 160) -> Iterable[Chunk]:
    for record in records:
        text = record.text
        if len(text) <= max_chars:
            yield record
            continue
        start = 0
        part = 1
        while start < len(text):
            end = min(start + max_chars, len(text))
            yield Chunk(text[start:end], record.metadata | {"part": part})
            if end == len(text):
                break
            start = max(0, end - overlap)
            part += 1


def excel_chunk(path: Path, sheet: str, row_start: int, row_end: int, rows: list[str]) -> Chunk:
    return Chunk(
        "\n".join(rows),
        base_metadata(path, "xlsx") | {"sheet": sheet, "row_start": row_start, "row_end": row_end},
    )


def base_metadata(path: Path, file_type: str) -> dict[str, str]:
    return {"file": path.name, "path": str(path), "type": file_type}


def format_cell(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def clean_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def make_chunk_id(chunk: Chunk) -> str:
    raw = f"{chunk.metadata}|{chunk.text[:120]}"
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()


if __name__ == "__main__":
    count = ingest_folder(os.environ.get("DOCS_DIR", "docs"), reset=True)
    print(f"Indexed {count} chunks.")
